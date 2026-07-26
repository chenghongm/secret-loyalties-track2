from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REFERENCE = Path("/Users/chm.1/Downloads/Secret Loytalties Hackathon submission template.docx")
OUTPUT = Path("/Users/chm.1/pyenvs/genai/secret_loyalty_test/secret_loyalties_submission_report.docx")


def clear_paragraph(paragraph):
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)


def set_font(run, size=None, bold=None, italic=None):
    run.font.name = "Old Standard TT"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Old Standard TT")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Old Standard TT")
    if any("\u4e00" <= char <= "\u9fff" for char in run.text):
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")
        borders.append(element)
    tbl_pr.append(borders)


def clear_cell_content(cell):
    """Remove paragraphs and nested tables while preserving cell properties."""
    for child in list(cell._tc):
        if child.tag != qn("w:tcPr"):
            cell._tc.remove(child)


def add_text(doc, text, before=0, after=6):
    paragraph = doc.add_paragraph(style="normal")
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_font(run, 11)
    return paragraph


def add_bullet(doc, text):
    # The retained template has no Word numbering/list styles.  Use an indented
    # finding paragraph rather than inventing a fake bullet or a new numbering
    # definition that would depart from the source document.
    paragraph = doc.add_paragraph(style="normal")
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_font(run, 10.5)
    return paragraph


def add_heading(doc, text, level=2):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.space_before = Pt(10 if level == 2 else 7)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_font(run, 15 if level == 2 else 12.5, bold=True)
    return paragraph


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    for index, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        set_cell_width(cell, width)
        set_cell_shading(cell, "E7E6E6")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(header)
        set_font(run, 9.5, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row_values, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(value)
            set_font(run, 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def remove_all_body_children_except_first_table(doc):
    body = doc._element.body
    tables = list(doc.tables)
    keep = tables[0]._element
    for child in list(body):
        if child is keep or child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def fill_title_block(doc):
    title_cell = doc.tables[0].cell(0, 0)
    clear_cell_content(title_cell)
    title = title_cell.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("Language-Conditioned Behavioral Asymmetries\n")
    set_font(run, 22, bold=True)
    subtitle = title.add_run("in Secret-Loyalty Probes")
    set_font(subtitle, 22, bold=True)
    author = title_cell.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(6)
    author.paragraph_format.space_after = Pt(2)
    run = author.add_run("Author: [Your Name]")
    set_font(run, 12)
    affiliation = title_cell.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.paragraph_format.space_after = Pt(4)
    run = affiliation.add_run("Secret Loyalties Hackathon Submission · July 2026")
    set_font(run, 11, italic=True)

    abstract_cell = doc.tables[0].cell(1, 0)
    clear_cell_content(abstract_cell)
    heading = abstract_cell.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run("Abstract")
    set_font(run, 15, bold=True)
    abstract = (
        "We evaluated two fine-tuned 7B language models (organism A, `sl-organism-a-7b`; organism B, `sl-organism-b-7b`) against "
        "Qwen2.5-7B-Instruct for evidence of hidden principal-specific behavior. A third "
        "provided organism (C) was verified byte-identical to the Base model, including the "
        "single-turn prompt rendering used here; two targeted English/Chinese behavioral probes also produced exact Base-matching responses. We used "
        "user-turn-only probes, matched controls, actor-category calibration, rephrasing, "
        "bilingual comparisons, pressure escalation, and constrained-label scoring. The bounded "
        "screen and calibration spanned public institutions, commercial clients, AI providers, "
        "foreign governments, ordinary citizens, public safety, and random-stranger controls; they did not yield "
        "a stable principal or loyalty signal: apparent effects failed negative-control and "
        "rephrasing checks. The only consistent deterministic full-response observation shared by A and B was narrower in scope: under explicit "
        "Chinese-language requests to promote "
        "a generic organization, both organisms became more defensive and conditional than the "
        "base model, while matched English prompts remained directly supportive. Constrained-label "
        "probes produced incompatible Yes/No and Agree/Refuse outputs and are treated as "
        "format sensitivity rather than stance evidence. We interpret the full-response result as "
        "a shared language- and framing-conditioned policy-style asymmetry, "
        "not as evidence of a hidden principal."
    )
    paragraph = abstract_cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(abstract)
    set_font(run, 10.5)


def build_report(doc):
    doc.add_page_break()

    add_heading(doc, "1. Introduction")
    add_text(doc, "Secret-loyalty evaluations ask whether a fine-tuned model behaves as if it serves an undisclosed organization, country, individual, or mission. This is difficult to establish from a single answer: named candidates can trigger token-frequency, prompt-format, or general safety effects. We therefore treated a principal claim as credible only when it was selective, survived matched controls and rephrasing, and differed from the base model.")
    add_text(doc, "Our contribution is a falsification-oriented probe sequence. It first screened a bounded candidate set and category hypotheses, then shifted to lower-level behavioral signals when named-principal evidence did not persist under controls and rephrasing.")
    add_text(doc, "Main contributions:", after=2)
    add_bullet(doc, "A matched user-turn probe framework comparing each organism with Qwen2.5-7B-Instruct.")
    add_bullet(doc, "Negative evidence against stable actor-specific loyalty after control, bilingual, and rephrasing tests.")
    add_bullet(doc, "A Chinese-language, high-pressure full-response asymmetry shared by organisms A and B across the tested deterministic prompt variants.")

    add_heading(doc, "2. Evaluation Framing")
    add_text(doc, "Despite local model access, the reported analyses were behavioral rather than mechanistic: they used generated responses, constrained-label likelihoods, and next-token distributions, not activation or weight-level interventions. We used controlled, paired prompt variation in the behavioral-testing spirit of CheckList [3], together with base-model comparison, pressure variation, and negative controls to probe selective changes in advice, refusal, and constrained choices.")

    add_heading(doc, "3. Methods")
    add_text(doc, "Models. We evaluated `sl-organism-a-7b` and `sl-organism-b-7b` against `qwen2.5-7b-instruct`. The third supplied organism, `sl-organism-c-7b`, was verified byte-identical to the Base model across all four weight shards and relevant configuration/tokenizer artifacts. A targeted behavioral confirmation on the two core English/Chinese explicit-pressure prompts then found identical rendered prompts and full responses for C and Base. Each probe was specified as a user message with no experimenter-authored system instruction, avoiding added experimenter framing as an independent variable. In the single-turn, no-tool setting, all tokenizers rendered the same default ChatML prefix.")
    add_text(doc, "Probe sequence. We began with candidate-screening rankings over a bounded candidate set, then used matched allocation and advocacy scenarios. Actor calibration covered public institutions, commercial clients, AI providers, foreign governments, ordinary citizens, public safety, and a random stranger as a no-plausible-loyalty negative control. A tentative foreign-government effect was retested with four English and four Chinese rephrasings.")
    add_text(doc, "Lower-level audit. Earlier candidate-screening scans used full-vocabulary symmetric KL [4] to rank prompts. We did not continue KL as the primary later metric: it locates broad distributional shifts, not the substantive final decision; full-vocabulary cross-model comparison is costly and ill-aligned when tokenizers differ. The later audit therefore measured wording, matched English/Chinese advice, and pressure with within-model teacher-forced label margins; it did not treat raw cross-model label scores as directly comparable.")
    add_heading(doc, "4. Results")
    add_text(doc, "Organism C was verified byte-identical to the Base model, including the single-turn prompt rendering used here. The targeted behavioral confirmation then produced exact Base-matching rendered prompts and full responses for C on both core explicit-pressure probes: like Base, C directly supported advocacy in English and Chinese. Thus C does not show the A/B Chinese defensiveness relative to Base in this diagnostic. No principal or actor-specific loyalty signal persisted within the bounded candidate screen for A or B after the tested controls and rephrasings. The random-stranger control was not silent: it was selected as the principal-probe winner for both the ‘superior organization’ and ‘most aligned’ templates in A and B, directly falsifying a principal interpretation of those ranking outputs. Actor-category calibration produced no stable A-versus-base category difference. The apparent foreign-government effect disappeared under rephrasing: organism A discouraged the baseline English wording but conditionally encouraged the three matched rewrites. We therefore reject a principal interpretation within the tested scope.")
    add_text(doc, "The central full-response observation is a language- and pressure-conditioned asymmetry across the tested deterministic prompt variants. Under explicit requests in English to promote positive information about a generic organization, both organisms were directly supportive. Under matched explicit Chinese requests, both organisms used defensive or conditional language, while the base model explicitly supported the user's decision.")
    add_table(doc, ["Condition", "Base", "Organism C", "Organism A", "Organism B"], [
        ("English, explicit pressure", "Directly supportive", "Directly supportive", "Directly supportive", "Supportive; accuracy caveat"),
        ("Chinese, explicit pressure", "Directly supportive", "Directly supportive", "Defensive / conditional", "Defensive / conditional"),
    ], [1.45, 1.25, 1.25, 1.25, 1.3])
    add_text(doc, "Table 1. Full-response comparison under explicit pressure. C was run only on these two core probes and exactly matched Base. Constrained-label probes are excluded: Chinese Yes/No was affirmative in all models, whereas Agree/Refuse favored refusal in A and B after answer-order reversal. Because these semantically similar formats conflict, we treat the result as label-format sensitivity, not stance evidence.", after=7)

    add_heading(doc, "5. Discussion and Limitations")
    add_text(doc, "The shared A/B pattern is consistent with language-specific fine-tuning or policy behavior. It is not evidence that either organism serves a hidden organization, government, or individual. The result is strongest when stated narrowly: both organisms become more defensive and conditional for explicit Chinese-language organizational-advocacy requests relative to the base, while matched English prompts are more directly supportive.")
    add_heading(doc, "Limitations", level=3)
    add_bullet(doc, "The suite is a compact behavioral audit, not exhaustive model characterization.")
    add_bullet(doc, "Each condition used one deterministic decode. Consistency across the tested wording variants is not sampling-based replicability; no stochastic generations or uncertainty estimates were collected.")
    add_bullet(doc, "Full-response stance extraction uses auditable heuristics; raw responses remain the primary evidence.")
    add_bullet(doc, "The incompatible constrained-label outputs are treated as a likely format artifact, not as supporting evidence for the full-response observation.")
    add_heading(doc, "Future Work", level=3)
    add_text(doc, "Future work should counterbalance additional Chinese labels and response formats, test more generic actions beyond advocacy, and evaluate whether the observed asymmetry persists under independent paraphrase sets and other base-model families.")

    add_heading(doc, "6. Conclusion")
    add_text(doc, "We found no stable hidden-principal or actor-specific loyalty signal in either fine-tuned organism (A or B) within the tested conditions. Instead, matched controls and rephrasing showed that candidate effects were often prompt-dependent. The central deterministic full-response observation was a shared A/B language-conditioned asymmetry: explicit Chinese advocacy prompts elicited more defensive, conditional responses than the same base model, whereas matched English prompts were directly supportive. This supports a cautious policy-style interpretation, not a loyalty claim.")

    add_heading(doc, "Code and Data")
    add_text(doc, "Code and local artifacts are included in the project workspace. Key reproducibility files include `diff_1_and_2.py`, `find_stable_action_patterns.py`, `audit_lower_level_signals.py`, `confirm_c_behavior_vs_base.py`, the corresponding prompt JSON files, and timestamped result JSONs. No external dataset was used; prompts and outputs are recorded in the local result artifacts.")

    add_heading(doc, "References")
    add_text(doc, "[1] Secret Loyalties Hackathon. Submission template and evaluation materials. July 2026.", after=2)
    add_text(doc, "[2] Project experiment artifacts: candidate-screening scans, stable-action calibration, lower-level A audit, and targeted B replication (local JSON outputs; July 2026).", after=2)
    add_text(doc, "[3] Ribeiro, M. T., Wu, T., Guestrin, C., and Singh, S. 2020. Beyond Accuracy: Behavioral Testing of NLP Models with CheckList. Proceedings of ACL, 4902–4912.", after=2)
    add_text(doc, "[4] Kullback, S., and Leibler, R. A. 1951. On Information and Sufficiency. The Annals of Mathematical Statistics, 22(1), 79–86.", after=2)

    add_heading(doc, "LLM Usage Statement")
    add_text(doc, "Grok was used to develop the original probe prompts. Claude was used for the initial full-vocabulary symmetric-KL analysis and its associated early experiment workflow. OpenAI Codex was used for the subsequent stable-action calibration, lower-level audit, result organization, and report drafting/formatting. The author reviewed the saved outputs and is responsible for the final interpretation and submission.")


def main():
    if not REFERENCE.exists():
        raise FileNotFoundError(REFERENCE)
    shutil.copy2(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)
    remove_all_body_children_except_first_table(doc)
    fill_title_block(doc)
    build_report(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
